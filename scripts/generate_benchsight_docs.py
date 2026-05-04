from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = REPO_ROOT / "output" / "doc"


def set_base_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_title_page(document: Document, title: str, subtitle: str) -> None:
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(22)

    subtitle_para = document.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(13)

    document.add_paragraph("")
    for line in [
        "Prepared for: ________________________________",
        "Course: _____________________________________",
        "Group Members: ______________________________",
        "Emails: _____________________________________",
        "Instructor / TA: _____________________________",
        "Submission Date: _____________________________",
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)

    document.add_section(WD_SECTION_START.NEW_PAGE)


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_workload_table(document: Document) -> None:
    document.add_heading("Workload Suite", level=1)
    document.add_paragraph(
        "BenchSight focuses on a compact but representative workload suite that covers both classical GPU patterns and modern deep learning operators."
    )

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Workload"
    header[1].text = "Category"
    header[2].text = "Implementations"
    header[3].text = "Purpose"

    rows = [
        ("Vector Add", "Classical GPU", "Baseline + Optimized", "Memory-bound calibration workload"),
        ("Reduction", "Classical GPU", "Baseline + Optimized", "Synchronization and speedup comparison"),
        ("GEMM", "Classical GPU", "Baseline + Optimized", "Compute-heavy matrix multiplication study"),
        ("Conv2D", "Deep Learning", "Baseline + Optimized", "Vision-oriented convolution benchmark"),
        ("Softmax", "Deep Learning", "Baseline + Optimized", "Transformer-style normalization benchmark"),
    ]

    for workload, category, implementations, purpose in rows:
        cells = table.add_row().cells
        cells[0].text = workload
        cells[1].text = category
        cells[2].text = implementations
        cells[3].text = purpose


def build_report_doc() -> None:
    document = Document()
    set_base_style(document)
    add_title_page(
        document,
        "BenchSight",
        "A GPU Benchmarking and Profiling Workflow for Custom CUDA Kernels",
    )

    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "BenchSight is a GPU benchmarking and profiling workflow designed to evaluate custom CUDA kernels under a repeatable measurement methodology. "
        "The project addresses a common workflow gap in performance engineering: benchmark experiments are frequently executed with ad hoc scripts, "
        "inconsistent timing practices, scattered outputs, and weak comparison workflows. BenchSight organizes this process into a clearer structure by "
        "combining custom CUDA workloads, warmup-plus-timed trial methodology, structured JSON result export, a Streamlit visualization dashboard, and "
        "Nsight-based profiler validation. The result is a project that is both technically meaningful and presentation-ready."
    )

    document.add_heading("1. Problem Overview", level=1)
    document.add_paragraph(
        "GPU optimization work is often evaluated through one-off scripts or isolated profiling sessions. In these workflows, warmup may be inconsistent, "
        "timing can be measured incorrectly, results are not stored in a structured form, and comparisons between implementations become difficult. "
        "As a result, performance claims are harder to validate and experimental conclusions become less reliable."
    )

    document.add_heading("2. Motivation and Importance", level=1)
    document.add_paragraph(
        "Reliable benchmark infrastructure matters because poor methodology leads to incorrect optimization decisions. If benchmark runs are not repeatable "
        "or comparable, then even promising-looking improvements may be misleading. BenchSight is motivated by the need for a structured workflow that makes "
        "it easier to measure, compare, and communicate GPU performance results with greater confidence."
    )

    document.add_heading("3. Current Approaches and Limitations", level=1)
    document.add_paragraph(
        "Current approaches typically fall into three categories: ad hoc scripts, standalone profiling tools, and large benchmark suites. Ad hoc scripts are "
        "easy to write but usually lack structure and repeatability. Profiling tools such as Nsight provide deep technical insight but are not designed to "
        "manage the full benchmark comparison workflow. Large benchmark suites are powerful but can be too complex for focused experiments and course-scale projects."
    )

    document.add_heading("4. Proposed Solution", level=1)
    document.add_paragraph(
        "BenchSight addresses this gap by combining custom CUDA workloads, consistent benchmark methodology, structured result export, visual comparison, and "
        "profiler-backed analysis. The current project scope uses a local Streamlit dashboard as the primary presentation surface. Benchmark runs are exported "
        "as JSON files, loaded into the dashboard, and visualized through latency, throughput, stability, and speedup views."
    )

    add_workload_table(document)

    document.add_heading("5. Project Flow", level=1)
    document.add_paragraph(
        "The project flow begins with a custom CUDA kernel implementation, either baseline or optimized. The kernel is executed using a benchmark harness that "
        "performs warmup iterations and timed trials. CUDA event timing is used to collect execution measurements, and NVTX ranges are used so the benchmark "
        "phases can be observed in Nsight. Each run is then exported as a structured JSON result file, which the Streamlit dashboard reads and visualizes. "
        "Profiler outputs from Nsight Systems and Nsight Compute are used alongside the dashboard to explain why one implementation outperforms another."
    )

    document.add_heading("6. Benchmark Methodology", level=1)
    add_bullet_list(
        document,
        [
            "Configure workload type and input parameters",
            "Execute warmup runs before timed trials",
            "Measure execution using CUDA events",
            "Compute average latency, p50, p95, throughput, and speedup",
            "Export each benchmark run as JSON",
            "Visualize the results in the Streamlit dashboard",
            "Validate execution behavior using Nsight Systems and Nsight Compute",
        ],
    )

    document.add_heading("7. Initial Results", level=1)
    document.add_paragraph(
        "The current project includes representative sample results across vector add, reduction, GEMM, Conv2D, and softmax. These results are presented in "
        "the dashboard as workload-level comparisons between baseline and optimized kernels. The dashboard highlights the fastest run, highest throughput, "
        "largest speedup, and stability trends across repeated timed runs."
    )
    document.add_paragraph(
        "Placeholder: Insert dashboard screenshots and measured benchmark figures here once the final experiments are complete."
    )

    document.add_heading("8. Why BenchSight Is Useful", level=1)
    document.add_paragraph(
        "BenchSight is not intended to replace Nsight. Instead, it complements profiler tools by acting as the evaluation and presentation layer around GPU "
        "benchmark experiments. Nsight explains why a kernel behaves a certain way. BenchSight measures, compares, and presents benchmark evidence across runs "
        "and implementations. This makes performance work easier to reproduce, easier to analyze, and easier to present."
    )

    document.add_heading("9. Limitations and Future Work", level=1)
    document.add_paragraph(
        "The current scope focuses on custom baseline and optimized CUDA kernels, a local JSON result workflow, and a single Streamlit dashboard. It does not "
        "currently include vendor-library references such as cuBLAS or cuDNN, automated profiler ingestion, or a full backend orchestration layer. Future work "
        "could add vendor-library comparison paths, additional workloads such as LayerNorm or attention, and tighter integration between benchmark outputs and profiler artifacts."
    )

    document.add_heading("10. Conclusion", level=1)
    document.add_paragraph(
        "BenchSight demonstrates that benchmarking methodology is itself a meaningful systems contribution. By combining custom CUDA kernels, structured timing, "
        "JSON result export, dashboard visualization, and profiler-backed interpretation, the project creates a cleaner and more credible workflow for GPU performance evaluation."
    )

    document.add_heading("Appendix: Final Fill-In Checklist", level=1)
    add_bullet_list(
        document,
        [
            "Insert course name, team names, and submission date",
            "Replace placeholder benchmark values with measured results",
            "Add final hardware and software configuration",
            "Insert dashboard screenshots",
            "Insert Nsight Systems and Nsight Compute screenshots",
            "Update workload-specific discussion with final observations",
        ],
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DIR / "BenchSight_Project_Report_Reframed.docx")


def build_presentation_doc() -> None:
    document = Document()
    set_base_style(document)
    add_title_page(
        document,
        "BenchSight Presentation Notes",
        "Slide-by-slide speaking content and question preparation",
    )

    slides = [
        (
            "Slide 1: Title",
            ["BenchSight", "A GPU Benchmarking and Profiling Workflow for Custom CUDA Kernels"],
            "Today I am presenting BenchSight, a GPU benchmarking and profiling workflow designed to evaluate custom CUDA kernels under a repeatable measurement methodology.",
        ),
        (
            "Slide 2: Problem Overview",
            [
                "Lack of standardized benchmarking methodology",
                "Inconsistent timing practices",
                "No structured result collection",
                "Difficult to compare and reproduce experiments",
            ],
            "The core problem is that GPU experiments are often run with ad hoc scripts. Timing practices vary, warmup may be skipped, and results are hard to compare fairly.",
        ),
        (
            "Slide 3: Motivation and Importance",
            [
                "Incorrect optimization decisions",
                "Reduced confidence in results",
                "Slower debugging and iteration",
                "Harder to validate improvements",
            ],
            "This matters because weak methodology leads to weak conclusions. BenchSight improves confidence by making performance measurement more structured and repeatable.",
        ),
        (
            "Slide 4: Current Approaches",
            [
                "Ad hoc benchmark scripts",
                "Standalone profiling tools",
                "Large benchmark suites",
            ],
            "Existing tools each solve part of the problem, but not the full benchmarking workflow. That is the gap BenchSight is trying to address.",
        ),
        (
            "Slide 5: Our Proposed Solution",
            [
                "Custom CUDA workloads",
                "Structured benchmark methodology",
                "JSON result export",
                "Streamlit dashboard",
                "Nsight-backed analysis",
            ],
            "BenchSight combines execution, measurement, comparison, and presentation into one coherent workflow.",
        ),
        (
            "Slide 6: Workload Suite",
            [
                "Vector Add",
                "Reduction",
                "GEMM",
                "Conv2D",
                "Softmax",
            ],
            "The workload suite covers both classical GPU behavior and deep learning operators while staying manageable for a focused project.",
        ),
        (
            "Slide 7: Project Workflow",
            [
                "Configure workload",
                "Warmup runs",
                "Timed CUDA trials",
                "Aggregate metrics",
                "Store as JSON",
                "Visualize in dashboard",
                "Validate with Nsight",
            ],
            "This slide shows the full project flow from kernel execution to presentation-ready evidence.",
        ),
        (
            "Slide 8: Initial Results",
            [
                "Baseline vs optimized latency",
                "Throughput comparison",
                "Speedup summary",
                "Profiler screenshots",
            ],
            "The dashboard and profiler views work together: the dashboard shows the result, and Nsight helps explain why it happened.",
        ),
        (
            "Slide 9: Why BenchSight Instead of Only Nsight?",
            [
                "Nsight is the profiler",
                "BenchSight is the benchmarking and comparison layer",
                "BenchSight standardizes methodology",
                "BenchSight improves presentation and reproducibility",
            ],
            "Nsight helps diagnose performance. BenchSight helps measure, compare, and communicate performance across implementations and runs.",
        ),
        (
            "Slide 10: Conclusion",
            [
                "Structured benchmark workflow",
                "Custom CUDA kernel comparison",
                "Profiler-backed validation",
                "Clear path for future expansion",
            ],
            "BenchSight turns GPU benchmarking into a more credible, reproducible, and presentation-ready workflow.",
        ),
    ]

    for heading, bullets, notes in slides:
        document.add_heading(heading, level=1)
        add_bullet_list(document, bullets)
        document.add_paragraph("Speaker Notes:", style="Intense Quote")
        document.add_paragraph(notes)

    document.add_heading("Potential Questions", level=1)
    add_bullet_list(
        document,
        [
            "Why not just use Nsight?",
            "Why is this better than an ad hoc script?",
            "Are you using your own kernels or vendor libraries?",
            "How do you ensure fair comparison?",
            "What would you add next if you had more time?",
        ],
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DIR / "BenchSight_Presentation_Notes.docx")


if __name__ == "__main__":
    build_report_doc()
    build_presentation_doc()
    print(f"Generated documents in {OUTPUT_DIR}")
